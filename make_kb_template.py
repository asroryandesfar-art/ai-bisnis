from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_default_font(document: Document, name: str = "Arial", size_pt: int = 12) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    # Ensure Latin/EastAsia font mapping
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.element.rPr.rFonts.set(qn("w:cs"), name)


def _set_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _title(document: Document, text: str, subtitle: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = document.add_paragraph(subtitle)
    p2.runs[0].font.size = Pt(12)
    p2.runs[0].italic = True


def _note_box(document: Document, label: str, body: str) -> None:
    t = document.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(body)
    document.add_paragraph()


def _qa_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Keyword"
    hdr[1].text = "Pertanyaan (contoh)"
    hdr[2].text = "Jawaban (versi final)"

    for k, q, a in rows:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = q
        row[2].text = a

    document.add_paragraph()


def build(path_out: str) -> None:
    doc = Document()
    _set_page(doc)
    _set_default_font(doc, "Arial", 12)

    today = date.today().isoformat()
    _title(
        doc,
        "BotNesia Knowledge Base (Template)",
        f"Versi: {today}  •  Bahasa: Indonesia  •  Isi dokumen ini lalu upload ke Dashboard → Documents",
    )

    _note_box(
        doc,
        "Cara pakai:",
        "Isi bagian-bagian di bawah dengan informasi bisnis kamu (SOP/FAQ). "
        "Gunakan bahasa yang jelas dan konsisten. Setelah selesai, export ke PDF (opsional) dan upload DOCX/TXT.",
    )

    doc.add_heading("1. Profil Bisnis", level=1)
    doc.add_paragraph("Nama bisnis: …")
    doc.add_paragraph("Jam operasional CS: …")
    doc.add_paragraph("Kontak CS (email/WA): …")
    doc.add_paragraph("Link kebijakan/terms (jika ada): …")

    doc.add_heading("2. Aturan Jawaban Bot", level=1)
    doc.add_paragraph("Tone: ramah, singkat, profesional.")
    doc.add_paragraph("Jika info tidak ada, minta detail yang relevan (nomor pesanan, email, tanggal).")
    doc.add_paragraph("Jika user marah/ancam legal, eskalasi ke human agent.")

    doc.add_heading("3. FAQ Utama", level=1)
    doc.add_heading("3.1 Login / Daftar", level=2)
    _qa_table(
        doc,
        [
            (
                "login",
                "Saya tidak bisa login, selalu gagal.",
                "Coba pastikan email dan password benar. Jika lupa password, gunakan fitur reset. "
                "Jika masih gagal, kirim screenshot error dan email yang dipakai.",
            ),
            (
                "daftar",
                "Daftar akun gagal / email sudah terdaftar.",
                "Jika email sudah terdaftar, silakan login atau gunakan fitur lupa password. "
                "Kalau masih error, kirim pesan errornya agar kami cek.",
            ),
        ],
    )

    doc.add_heading("3.2 Pengiriman", level=2)
    _qa_table(
        doc,
        [
            (
                "pengiriman",
                "Pesanan saya belum sampai, bagaimana cek statusnya?",
                "Kirim nomor pesanan atau nomor resi. Kami akan cek status (diproses/dikirim/tertahan) "
                "dan bantu tindak lanjut sesuai kondisi.",
            ),
            (
                "estimasi",
                "Estimasi pengiriman berapa hari?",
                "Estimasi pengiriman: … hari kerja (tergantung lokasi). Jika sudah lewat estimasi, kirim nomor pesanan untuk dicek.",
            ),
        ],
    )

    doc.add_heading("3.3 Refund / Retur", level=2)
    _qa_table(
        doc,
        [
            (
                "refund",
                "Saya mau refund, prosedurnya bagaimana?",
                "Mohon kirim nomor pesanan, tanggal transaksi, metode pembayaran, dan alasan refund. "
                "Estimasi proses refund: … hari kerja setelah disetujui.",
            ),
            (
                "retur",
                "Barang rusak, bisa retur?",
                "Bisa. Mohon kirim nomor pesanan + foto/video kondisi barang. Kami cek dan informasikan langkah retur/penggantian.",
            ),
        ],
    )

    doc.add_heading("3.4 Harga / Paket", level=2)
    _qa_table(
        doc,
        [
            (
                "harga",
                "Harga paket dan fiturnya apa saja?",
                "Paket tersedia: … (isi detail). Kalau kamu sebutkan kebutuhan (jumlah bot & chat/bulan), kami bantu rekomendasikan paket.",
            ),
        ],
    )

    doc.add_heading("4. SOP Eskalasi (Human Agent)", level=1)
    doc.add_paragraph("Eskalasi jika:")
    doc.add_paragraph("• User meminta bicara manusia / admin.", style="List Bullet")
    doc.add_paragraph("• Ada ancaman legal / publik / emosi sangat negatif.", style="List Bullet")
    doc.add_paragraph("• Kendala teknis berulang / error server.", style="List Bullet")
    doc.add_paragraph("Template balasan eskalasi:")
    doc.add_paragraph(
        "“Baik, saya bantu hubungkan ke tim kami agar ditangani lebih cepat. Mohon tunggu sebentar ya.”"
    )

    doc.add_heading("5. Lampiran (Opsional)", level=1)
    doc.add_paragraph("Tambah kebijakan lengkap, daftar produk, atau SOP internal di sini.")

    doc.save(path_out)


if __name__ == "__main__":
    build("BotNesia_Knowledge_Base_Template.docx")

